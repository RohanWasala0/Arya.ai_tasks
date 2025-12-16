import docx
import aspose.words as aw
from aspose.words import NodeType

def read_tables_data(filename):
    document = docx.Document(filename)
    all_tables_data = []
    for table in document.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            print(row_data)
            table_data.append(row_data)
        all_tables_data.append(table_data)
    return all_tables_data

def read_table_with_coor(filename):
    # Load the document
    doc = aw.Document(filename, load_options= None)

    # Get all tables in the document
    tables = doc.get_child_nodes(node_type= aw.NodeType.TABLE, is_deep= True)

    table_info_list = []

    for idx, table in enumerate(tables, 1):
        table_node = table.as_table()
        
        # Get the layout collector to access positioning information
        layout_collector = aw.layout.LayoutCollector(doc)
        layout_enumerator = aw.layout.LayoutEnumerator(doc)
        
        # Get the first cell to determine table position
        first_cell = table_node.first_row.first_cell
        
        # Get page number
        page_number = layout_collector.get_start_page_index(table_node)
        
        # Move enumerator to the table
        layout_enumerator.set_current(layout_collector.get_entity(table_node))
        
        # Get rectangle (coordinates) of the table
        rect = layout_enumerator.rectangle
        
        table_info = {
            'table_index': idx,
            'page_number': page_number,
            'coordinates': {
                'x': rect.x,
                'y': rect.y,
                'width': rect.width,
                'height': rect.height,
                'left': rect.x,
                'top': rect.y,
                'right': rect.x + rect.width,
                'bottom': rect.y + rect.height
            },
            'row_count': table_node.rows.count,
            'column_count': table_node.first_row.cells.count if table_node.rows.count > 0 else 0
        }
        
        table_info_list.append(table_info)
        
        # Print information
        print(f"\n{'='*60}")
        print(f"Table {idx}:")
        print(f"  Page: {page_number}")
        print(f"  Position (X, Y): ({rect.x:.2f}, {rect.y:.2f})")
        print(f"  Size (W x H): {rect.width:.2f} x {rect.height:.2f}")
        print(f"  Bounds: Left={rect.x:.2f}, Top={rect.y:.2f}, "
              f"Right={rect.x + rect.width:.2f}, Bottom={rect.y + rect.height:.2f}")
        print(f"  Rows: {table_node.rows.count}, Columns: {table_node.first_row.cells.count if table_node.rows.count > 0 else 0}")
    
    return table_info_list

"""
Simple DOCX to PDF coordinate extraction using pdfplumber
Install: pip install pdfplumber docx2pdf
"""
import pdfplumber
from docx2pdf import convert

# Convert DOCX to PDF
pdf_file = "./bank_statment.pdf"

# Extract table coordinates
with pdfplumber.open(pdf_file) as pdf:
    for page_num, page in enumerate(pdf.pages, 1):
        tables = page.find_tables()
        
        for idx, table in enumerate(tables, 1):
            bbox = table.bbox  # (x0, top, x1, bottom)
            print(f"\nTable {idx} on Page {page_num}:")
            print(f"  Position: x={bbox[0]:.2f}, y={bbox[1]:.2f}")
            print(f"  Size: width={bbox[2]-bbox[0]:.2f}, height={bbox[3]-bbox[1]:.2f}")
            print(f"  Bounds: left={bbox[0]:.2f}, top={bbox[1]:.2f}, right={bbox[2]:.2f}, bottom={bbox[3]:.2f}")

# Process the list of lists returned by this function, potentially using pandas
# tables_data = read_tables_data('./Resources/bank_statment.docx')
# read_tables_data("./Resources/bank_statment.docx")

